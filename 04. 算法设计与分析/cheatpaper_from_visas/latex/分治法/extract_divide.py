"""
从分治法Word文档中提取题目
"""

import os
import re
from pathlib import Path
from docx import Document

# ============ 配置区 ============
SOURCE_FOLDER = './docs'          # Word文档所在文件夹
OUTPUT_FOLDER = './problems_txt'  # 输出文件夹
# =================================

def is_new_problem(text):
    """
    判断是否为新题目的开始
    根据分治法文档的格式调整
    """
    text = text.strip()

    # 自动检测模式
    patterns = [
        r'^大题\d+',                # 大题17、大题18等
        r'^题目\d+',                # 题目1、题目2等
        r'^问题\d+',                # 问题1、问题2等
        r'^\d+\.[\u4e00-\u9fa5]',   # "1. 题目"这种格式（数字+点+中文）
        r'^【\d+】',                # 【1】、【2】
        r'^###',                     # 手动标记 ###
    ]

    for pattern in patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True

    return False


def extract_text_from_docs():
    """从所有Word文档中提取题目"""

    # 创建输出文件夹
    Path(OUTPUT_FOLDER).mkdir(exist_ok=True)

    # 获取所有docx文件并排序
    doc_files = sorted([f for f in os.listdir(SOURCE_FOLDER) if f.endswith('.docx')])

    if not doc_files:
        print(f"在 {SOURCE_FOLDER} 文件夹中未找到 .docx 文件")
        return

    print(f"找到 {len(doc_files)} 个Word文档")
    print("=" * 60)

    problem_count = 1
    current_problem_lines = []
    current_title = ""

    for filename in doc_files:
        filepath = os.path.join(SOURCE_FOLDER, filename)
        print(f"正在处理: {filename}")

        try:
            doc = Document(filepath)

            for para in doc.paragraphs:
                text = para.text.strip()

                # 跳过空行
                if not text:
                    continue

                # 检测是否为新题目
                if is_new_problem(text):
                    # 保存上一个题目
                    if current_problem_lines:
                        save_problem(problem_count, current_title, current_problem_lines)
                        problem_count += 1

                    # 开始新题目
                    current_title = text
                    current_problem_lines = [text]
                else:
                    # 追加到当前题目
                    current_problem_lines.append(text)

            # 处理文档中的表格（有些题目可能在表格里）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text.strip()
                            if text and is_new_problem(text):
                                if current_problem_lines:
                                    save_problem(problem_count, current_title, current_problem_lines)
                                    problem_count += 1
                                current_title = text
                                current_problem_lines = [text]
                            elif text:
                                current_problem_lines.append(text)

        except Exception as e:
            print(f"处理 {filename} 时出错: {e}")
            continue

    # 保存最后一题
    if current_problem_lines:
        save_problem(problem_count, current_title, current_problem_lines)
        problem_count += 1

    print("=" * 60)
    print(f"提取完成！共生成 {problem_count - 1} 个题目文件")
    print(f"保存位置: {OUTPUT_FOLDER}/")


def save_problem(idx, title, lines):
    """保存单个题目到txt文件"""

    content = "\n".join(lines)

    # 过滤太短的内容（可能是误识别）
    if len(content) < 30:
        print(f"跳过过短内容: {title[:30]}...")
        return

    filename = f"{idx:02d}.txt"
    filepath = os.path.join(OUTPUT_FOLDER, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)

    print(f"  已保存: {filename} - {title[:40]}")


if __name__ == '__main__':
    print("=" * 60)
    print("分治法题目提取工具")
    print("=" * 60)
    print()

    # 检查源文件夹是否存在
    if not os.path.exists(SOURCE_FOLDER):
        print(f"源文件夹不存在: {SOURCE_FOLDER}")
        print(f"请创建此文件夹并放入Word文档")
    else:
        extract_text_from_docs()

    print()
    print("提示:")
    print(f"  - 检查 {OUTPUT_FOLDER}/ 文件夹中的提取结果")
    print(f"  - 如有识别错误，可在Word中手动添加 ### 标记题目开头")
    print(f"  - 修改脚本中的 is_new_problem() 函数可调整识别规则")
